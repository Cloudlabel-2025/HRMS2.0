from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('docs/user-guides')

ROLES = {
    'Super Admin': [
        ('Dashboard', 'Review company-wide highlights, pending items, and quick links.'),
        ('Employees', 'Create and maintain employee records, view the directory, and review individual profiles.'),
        ('Recruitment', 'Create job openings, add applicants, and move candidates through hiring stages.'),
        ('Attendance', 'Review attendance, correct records when appropriate, and handle regularization requests.'),
        ('Absence', 'Review absence information and employee requests.'),
        ('Leave', 'Review and decide leave requests; view balances and leave history.'),
        ('Tasks & Projects', 'Create, assign, follow up on, and update work items and projects.'),
        ('Monitoring', 'Review work and activity information for operational follow-up.'),
        ('Payroll', 'Prepare, review, approve, and finalize payroll; view payslips.'),
        ('Finance', 'Manage clients, expenses, budgets, and accounting records.'),
        ('Invoices', 'Create and track customer invoices.'),
        ('Inventory', 'Maintain company inventory and assigned items.'),
        ('Performance', 'Set goals and manage reviews across the organization.'),
        ('Documents', 'Upload, organize, and share HR and work documents.'),
        ('My Profile', 'Review and update your own profile and submit personal requests.'),
        ('Core HR', 'Maintain workforce profiles, lifecycle information, and HR requests.'),
        ('Announcements', 'Publish and manage company announcements.'),
        ('Calendar', 'View events and manage holidays or company calendar items.'),
        ('Reports', 'Run and review organization reports.'),
        ('SME Portal', 'Maintain subject-matter expert records and assignments.'),
        ('Settings', 'Maintain working settings, leave policies, and related master data.'),
        ('Audit Logs', 'Review the system activity history.'),
    ],
    'Admin': [
        ('Dashboard', 'Review company-wide highlights, pending items, and quick links.'),
        ('Employees', 'Create and maintain employee records, view the directory, and review individual profiles.'),
        ('Recruitment', 'Create job openings, add applicants, and move candidates through hiring stages.'),
        ('Attendance', 'Review attendance, correct records when appropriate, and handle regularization requests.'),
        ('Absence', 'Review absence information and employee requests.'),
        ('Leave', 'Review and decide leave requests; view balances and leave history.'),
        ('Tasks & Projects', 'Create, assign, follow up on, and update work items and projects.'),
        ('Monitoring', 'Review work and activity information for operational follow-up.'),
        ('Payroll', 'Prepare and review payroll; view payslips available to you.'),
        ('Finance', 'Work with the finance records made available to your role.'),
        ('Invoices', 'Work with the invoice records made available to your role.'),
        ('Inventory', 'Maintain company inventory and assigned items.'),
        ('Performance', 'Set goals and manage reviews across the organization.'),
        ('Documents', 'Upload, organize, and share HR and work documents.'),
        ('My Profile', 'Review and update your own profile and submit personal requests.'),
        ('Core HR', 'Maintain workforce profiles, lifecycle information, and HR requests.'),
        ('Announcements', 'Publish and manage company announcements.'),
        ('Calendar', 'View events and manage holidays or company calendar items.'),
        ('Reports', 'Run and review organization reports.'),
        ('Settings', 'Maintain the settings and leave policies made available to you.'),
    ],
    'Recruiter': [
        ('Dashboard', 'Review your work summary and hiring-related quick links.'),
        ('Employees', 'Search the employee directory and review employee details.'),
        ('Recruitment', 'Create job openings, add applicants, schedule the hiring flow, and update candidate stages.'),
        ('Performance', 'Review the performance information made available to you.'),
        ('Documents', 'Use the documents made available to you for hiring and employee work.'),
        ('My Profile', 'Review and update your own profile and submit personal requests.'),
        ('Core HR', 'View workforce records and HR information needed for recruitment.'),
        ('Calendar', 'View your calendar and create personal calendar items where available.'),
        ('Reports', 'Review the reports made available to you.'),
    ],
    'Team Admin': [
        ('Dashboard', 'Review your team summary, pending items, and quick links.'),
        ('Employees', 'View and manage the people in your team.'),
        ('Attendance', 'Review team attendance and handle team-level follow-up.'),
        ('Absence', 'Review absence information for your team.'),
        ('Leave', 'Review team leave requests and their status.'),
        ('Tasks & Projects', 'Create and assign work for your team; follow progress and update work details.'),
        ('Monitoring', 'Review team work and activity information.'),
        ('Performance', 'Set goals and review performance for your team.'),
        ('Documents', 'Use and manage documents for your team.'),
        ('My Profile', 'Review and update your own profile and submit personal requests.'),
        ('Core HR', 'View and manage the HR records available for your team.'),
        ('Announcements', 'Publish and review announcements for your team.'),
        ('Calendar', 'View and manage team calendar items.'),
        ('Reports', 'Run reports for your team.'),
        ('Inventory', 'Review and manage inventory items available to your team.'),
    ],
    'Team Lead': [
        ('Dashboard', 'Review your team summary, pending items, and quick links.'),
        ('Employees', 'View the employees you lead and their relevant details.'),
        ('Attendance', 'Review team attendance and handle team-level follow-up.'),
        ('Absence', 'Review absence information for your team.'),
        ('Leave', 'Review team leave requests and their status.'),
        ('Tasks & Projects', 'Assign work to your team, monitor progress, and update work details.'),
        ('Monitoring', 'Review team work and activity information.'),
        ('Performance', 'Set goals and review performance for your team.'),
        ('Documents', 'Use and manage documents for your team.'),
        ('My Profile', 'Review and update your own profile and submit personal requests.'),
        ('Core HR', 'View and manage the HR records available for your team.'),
        ('Announcements', 'Publish and review announcements for your team.'),
        ('Calendar', 'View and manage team calendar items.'),
        ('Reports', 'Run reports for your team.'),
        ('Inventory', 'Review and manage inventory items available to your team.'),
    ],
    'Employee': [
        ('Dashboard', 'Review your personal work summary, reminders, and quick links.'),
        ('Employees', 'Find employee information that is available to you.'),
        ('Attendance', 'Record your daily attendance, check history, and submit correction requests if needed.'),
        ('Absence', 'Review your absence information.'),
        ('Leave', 'Check your leave balance, request leave, and track the decision.'),
        ('Tasks & Projects', 'View assigned work, update progress, and work with assigned projects.'),
        ('Payroll', 'View your payroll and payslip information when it is published.'),
        ('Performance', 'Set or update your own goals and review your feedback.'),
        ('Documents', 'View the documents shared with you and your own documents.'),
        ('My Profile', 'Keep your profile information current and submit personal service requests.'),
        ('Announcements', 'Read company and team announcements.'),
        ('Calendar', 'View company and personal calendar events.'),
        ('Reports', 'View your personal reports where available.'),
    ],
    'Intern': [
        ('Dashboard', 'Review your personal work summary and quick links.'),
        ('Attendance', 'Record your daily attendance, check history, and submit correction requests if needed.'),
        ('Absence', 'Review your absence information.'),
        ('Leave', 'Request eligible leave and track its status.'),
        ('Tasks & Projects', 'View assigned work and update your progress.'),
        ('Performance', 'Review the performance information made available to you.'),
        ('Documents', 'View documents shared with you.'),
        ('My Profile', 'Review and update your own profile and submit personal requests.'),
        ('Announcements', 'Read company and team announcements.'),
        ('Calendar', 'View company and personal calendar events.'),
    ],
    'SME': [
        ('Dashboard', 'Review your personal work summary and quick links.'),
        ('Attendance', 'Record your daily attendance and review your attendance history.'),
        ('Leave', 'Check your leave balance, request leave, and track the decision.'),
        ('Tasks & Projects', 'View assigned work and projects, then update progress.'),
        ('Payroll', 'View your own payroll information when it is published.'),
        ('My Profile', 'Review and update your own profile and submit personal requests.'),
        ('Calendar', 'View company and personal calendar events.'),
    ],
}

SPECIAL_STEPS = {
    'Attendance': ['Open Attendance from the left menu.', 'Use the main action on the page to record your time or choose the relevant attendance option.', 'Check the date and status shown in the history area.', 'If a correction is needed, choose the regularization or correction option and submit the details.'],
    'Leave': ['Open Leave from the left menu.', 'Review your available leave balance and the leave types shown.', 'Choose New Leave Request, enter the dates and reason, then submit.', 'Return to Leave later to see whether the request is pending, approved, on hold, or declined.'],
    'Tasks & Projects': ['Open Tasks & Projects from the left menu.', 'Choose a task or project from the list.', 'Read the description, due date, and current status.', 'Update your status or progress, then save. Managers can also create and assign work.'],
    'My Profile': ['Open My Profile from the left menu.', 'Review the information shown in the profile sections.', 'Choose Edit or the relevant request option when you need to change permitted details or raise a request.', 'Save your update and check the request status on a later visit.'],
    'Calendar': ['Open Calendar from the left menu.', 'Use the month, week, or day controls to find the date you need.', 'Select an event to read its details.', 'Where an Add Event option is available, enter the details and save.'],
    'Documents': ['Open Documents from the left menu.', 'Use search or category filters to find a document.', 'Select the document name to open or download it.', 'If you see an upload option, choose the file, add the requested details, and save.'],
    'Performance': ['Open Performance from the left menu.', 'Choose Goals or Reviews, depending on what you need to do.', 'Open an existing item or choose the add option if it is available.', 'Enter your update or review details and save.'],
    'Payroll': ['Open Payroll from the left menu.', 'Choose the relevant month or payroll record.', 'Review the amounts and status shown.', 'Use the payslip view or download option when it is available.'],
    'Employees': ['Open Employees from the left menu.', 'Use search and filters to find the person or group you need.', 'Select a name to open the profile.', 'Use the available actions to view, update, or manage information within your responsibility.'],
    'Recruitment': ['Open Recruitment from the left menu.', 'Choose Jobs to manage openings or Applicants to manage candidates.', 'Use the add option to create a new record, or select an existing record to update it.', 'Move the applicant to the appropriate stage and save your changes.'],
    'Dashboard': ['Open Dashboard from the left menu.', 'Review the summary cards and any highlighted pending items.', 'Select a card or quick link to go straight to that area.', 'Use this page as your starting point each day.'],
}

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn('w:tcW'))
    if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'), 'true'); trPr.append(el)

def add_footer(section, role):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f'HRMS User Guide | {role}'); r.font.name = 'Calibri'; r.font.size = Pt(8); r.font.color.rgb = RGBColor(100, 116, 139)

def add_title(doc, role):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('HRMS'); r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(30, 64, 175)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    r = p.add_run(f'{role} User Guide'); r.bold = True; r.font.name = 'Calibri'; r.font.size = Pt(25); r.font.color.rgb = RGBColor(15, 23, 42)
    p = doc.add_paragraph('A simple guide to the screens and actions available in your daily work.')
    p.paragraph_format.space_after = Pt(16); p.runs[0].font.color.rgb = RGBColor(71, 85, 105)

def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(4); p.add_run(step)

def generic_steps(module, audience):
    verb = 'review' if audience in ['Employee', 'Intern', 'SME'] else 'manage'
    return [f'Open {module} from the left menu.', f'Use the page filters, tabs, or search box to find the information you need.', f'Select the relevant item to {verb} its details.', 'Choose Save, Submit, or Update after making a change.']

def build(role, modules):
    doc = Document(); sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(.492)
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri'); normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [('Heading 1',16,'2E74B5',18,10),('Heading 2',13,'2E74B5',14,7),('Heading 3',12,'1F4D78',10,5)]:
        st=doc.styles[name]; st.font.name='Calibri'; st._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); st._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after)
    add_footer(sec, role); add_title(doc, role)
    doc.add_heading('How to use this guide', level=1)
    doc.add_paragraph('This guide uses the same menu names that appear in HRMS. Start from the left menu, then follow the short steps under the module you need. Your screen may show fewer actions than a colleague’s screen; use the options that appear for you.')
    doc.add_heading('Your module map', level=1)
    table = doc.add_table(rows=1, cols=2); table.alignment=WD_TABLE_ALIGNMENT.LEFT; table.style='Table Grid'; table.autofit=False
    widths=[2700,6660]
    for cell, text, width in zip(table.rows[0].cells, ['Menu item', 'What it helps you do'], widths):
        cell.text=text; shade(cell,'E8EEF5'); set_cell_width(cell,width); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs: run.bold=True
    set_repeat_table_header(table.rows[0])
    for module, purpose in modules:
        cells=table.add_row().cells
        for cell, text, width in zip(cells,[module,purpose],widths):
            cell.text=text; set_cell_width(cell,width); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_heading('Everyday navigation', level=1)
    doc.add_paragraph('The left menu is the main way to move around HRMS. The page title tells you where you are. Use tabs, filters, search, and the action buttons on each page to narrow down information or start a task. Use Back in your browser only when a page does not provide a clear return option.')
    for module, purpose in modules:
        doc.add_heading(module, level=1)
        doc.add_paragraph(purpose)
        doc.add_heading('Navigation path', level=2)
        doc.add_paragraph(f'Left menu > {module}')
        doc.add_heading('What to do', level=2)
        add_steps(doc, SPECIAL_STEPS.get(module, generic_steps(module, role)))
        if module in ['Leave','Attendance','Tasks & Projects','Payroll','Performance']:
            doc.add_heading('Helpful tip', level=2)
            tips = {'Leave':'Keep the request dates and reason clear so the reviewer has the information needed to decide.', 'Attendance':'Check the date before submitting a correction; it makes later review easier.', 'Tasks & Projects':'Update the status whenever your work changes so the team sees an accurate picture.', 'Payroll':'If an amount looks unexpected, note the month and payslip details before raising a question with the responsible team.', 'Performance':'Use specific examples when describing progress or feedback.'}
            doc.add_paragraph(tips[module])
    doc.add_heading('When you need help', level=1)
    doc.add_paragraph('If a required option is not visible or a record needs correction, contact your HR or team contact with the page name, the item name, and a short description of what you need. Do not share personal or confidential details in public announcements or comments.')
    filename = role.lower().replace(' ', '-') + '-user-guide.docx'
    doc.save(OUT / filename)

OUT.mkdir(parents=True, exist_ok=True)
for role, modules in ROLES.items(): build(role, modules)
print(f'Created {len(ROLES)} guides in {OUT.resolve()}')
